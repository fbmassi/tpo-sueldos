"""Documento de defensa — con feature engineering justificado, roles y por que RF."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT=Path(__file__).resolve().parent.parent
EDA=ROOT/"data/processed/eda"; MOD=ROOT/"data/processed/modelos"; PRE=ROOT/"data/processed/presentacion"
OUT=ROOT/"presentacion/Defensa_TPO.docx"
NAVY=RGBColor(0x0A,0x0A,0x0A); TEAL=RGBColor(0x33,0x33,0x33)

doc=Document(); doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(11)

def h(txt,level=1):
    p=doc.add_heading(txt,level=level)
    for r in p.runs: r.font.color.rgb=NAVY if level<=1 else TEAL
    return p
def para(txt,bold=False,italic=False,size=11,align=None,color=None):
    p=doc.add_paragraph(); r=p.add_run(txt); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    if align: p.alignment=align
    return p
def bullet(lead,rest=""):
    p=doc.add_paragraph(style="List Bullet")
    if rest: r=p.add_run(lead); r.bold=True; p.add_run(rest)
    else: p.add_run(lead)
    return p
def img(path,w=6.0,cap=None):
    doc.add_picture(str(path),width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if cap: para(cap,italic=True,size=9,align=WD_ALIGN_PARAGRAPH.CENTER)
def table(headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Light Grid Accent 1"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,hd in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(hd); r.bold=True; r.font.size=Pt(10)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""; r=cells[i].paragraphs[0].add_run(str(v)); r.font.size=Pt(10)
    return t

# PORTADA
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("Mercado Laboral Tech Argentino"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=NAVY
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("Análisis y predicción de salarios del sector IT"); r.font.size=Pt(16); r.font.color.rgb=TEAL
doc.add_paragraph()
para("Trabajo Práctico Obligatorio — Minería de Datos",bold=True,size=12,align=WD_ALIGN_PARAGRAPH.CENTER)
para("Documento de defensa oral",italic=True,size=12,align=WD_ALIGN_PARAGRAPH.CENTER)
para("Base: encuestas de Sysarmy 2022–2025 + contexto macroeconómico",size=11,align=WD_ALIGN_PARAGRAPH.CENTER)
para("Integrantes: ____________________     |     Fecha: ____________",size=11,align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# RESUMEN
h("Resumen ejecutivo",1)
para("Este trabajo integra seis ediciones de la encuesta salarial de Sysarmy (2022–2025) con ocho "
     "series macroeconómicas para construir un dataset único, limpio y comparable en el tiempo de "
     "31.088 profesionales del sector tecnológico argentino. Todos los salarios se expresan en "
     "términos reales (pesos y dólares constantes de mayo 2026), neutralizando una inflación "
     "acumulada superior al 900% y la fuerte variación del tipo de cambio del período.")
para("Sobre esa base se desarrollaron dos modelos: un Random Forest (supervisado) que estima el "
     "salario a partir del perfil, y una segmentación con K-Means (no supervisado) que descubre los "
     "arquetipos del mercado. Los resultados se exponen en una app web (Streamlit), en versiones "
     "español/pesos e inglés/dólares.")
para("Hallazgo central: el perfil explica alrededor del 27% de la variación salarial —la experiencia "
     "y el rol de liderazgo son dominantes— y, contra la intuición, cobrar en dólares no aportó "
     "mayor estabilidad en el período.",italic=True)

# EQUIPO Y ESTRATEGIA
h("Equipo, roles y estrategia",1)
para("El equipo se organizó por roles complementarios, bajo una metodología iterativa inspirada en "
     "CRISP-DM (comprensión del negocio → datos → preparación → modelado → evaluación → despliegue), "
     "con control de versiones en Git y ramas por funcionalidad.")
table(["Rol","Responsabilidades","Integrante"],[
    ["Datos / ETL","Descarga e integración de 9 fuentes, limpieza y unificación","____________"],
    ["Feature engineering / Modelado","Variables, Random Forest, K-Means y evaluación","____________"],
    ["EDA / Visualización","Análisis exploratorio, gráficos y storytelling","____________"],
    ["App / Documentación","App Streamlit, informe y presentación","____________"],
])

# 1 DOMINIO
h("1. Dominio y problemática",1)
para("El sector tecnológico argentino es dinámico, con alta rotación y exposición internacional. Sus "
     "salarios están atados a una macro muy volátil, lo que genera un problema concreto: no existe "
     "información integrada y comparable en el tiempo. Cada edición está en pesos de su momento, con "
     "un nivel de precios y un tipo de cambio distintos, por lo que comparar a valor nominal carece "
     "de sentido.")
bullet("Fragmentación de las fuentes ", "(encuesta + macro dispersa).")
bullet("Falta de una medida de salario real ", "comparable entre períodos.")
bullet("Ausencia de una herramienta ", "que traduzca el perfil en una estimación de mercado.")

# 2 HIPOTESIS
h("2. Hipótesis",1)
para("Hipótesis central",bold=True,color=TEAL)
para("El sueldo real de un profesional tech está determinado por su perfil laboral (rol, "
     "experiencia) y su contexto (provincia, modalidad, tamaño de empresa), una vez expresado en "
     "moneda constante para que las ediciones sean comparables entre sí.",italic=True)
para("Hipótesis desagregadas",bold=True,color=TEAL)
table(["#","Hipótesis"],[
    ["H1","El seniority/experiencia es la variable más predictiva."],
    ["H2","La experiencia suma con rendimientos decrecientes."],
    ["H3","Ciertas tecnologías tienen premium salarial."],
    ["H4","CABA paga más que el interior, pero menos de lo que se cree."],
    ["H5","Quienes cobran en dólares tienen poder de compra más estable."],
    ["H6","Cuando el ITCRM sube, suben los sueldos en USD."],
])

# 3 PROPUESTA
h("3. Propuesta y valor para el negocio",1)
para("Se construyó un pipeline reproducible que unifica las seis ediciones con ocho series macro, "
     "produce un dataset analítico único y alimenta modelos predictivos, de segmentación y una app.")
para("Valor para la gerencia:",bold=True)
bullet("Benchmarking salarial objetivo para ofertas y retención.")
bullet("Identificación de brechas y de los factores que mueven el sueldo.")
bullet("Lectura del impacto macro (dólar y tipo de cambio real) sobre el costo del talento.")
bullet("Herramienta viva que se reentrena con cada nueva edición.")

# 4 PIPELINE
h("4. Arquitectura del pipeline y fuentes",1)
para("Flujo: (1) descarga, (2) limpieza y unificación, (3) ajuste a valores reales con base común "
     "(mayo 2026) y (4) dataset final más una tabla de contexto macro vinculada por fecha de edición.")
para("Fuentes integradas (9):",bold=True)
table(["Fuente","Aporte"],[
    ["Sysarmy (6 ediciones)","Salarios y perfiles — 31.088 respuestas (2022–2025)"],
    ["IPC INDEC","Inflación argentina (deflactor de pesos)"],
    ["Dólar MEP (Blue, bluelytics)","Tipo de cambio histórico"],
    ["US CPI (FRED)","Inflación de EE.UU. (deflactor de dólares)"],
    ["CBT INDEC","Canasta básica total"],
    ["RIPTE","Salario formal promedio (benchmark)"],
    ["Big Mac Index","Paridad de poder adquisitivo"],
    ["ITCRM (BCRA)","Tipo de cambio real multilateral"],
    ["Stack Overflow","Benchmark internacional (separado)"],
])
para("Reproducibilidad: los CSV crudos no viajan con el repositorio (el de Stack Overflow pesa "
     "140 MB, por encima del límite de GitHub), pero cada fuente tiene un gemelo .parquet liviano "
     "versionado en data/raw/ (140 MB → 11 MB). El pipeline los usa automáticamente si el CSV no "
     "está, de modo que cualquier integrante reproduce todo el trabajo con un simple clon del "
     "repositorio.",italic=True)

# 5 LIMPIEZA
h("5. Limpieza de datos y dificultades resueltas",1)
para("La principal complejidad no fue limpiar datos sucios —los campos son en su mayoría "
     "desplegables controlados— sino integrar y hacer comparables fuentes heterogéneas. De 32.309 "
     "filas crudas se llegó a 31.088 finales (sólo 3,8% descartado). Decisiones clave:")
bullet("Armonización de esquemas: ","cada edición tenía 43–56 columnas con nombres distintos; se "
       "mapearon al mismo esquema por coincidencia de tokens.")
bullet("Parseo y tipado: ","salario de texto a número, fechas, normalización de texto.")
bullet("Valores faltantes: ","eliminación en columnas críticas (provincia, salario) e imputación por "
       "mediana/moda en el resto.")
bullet("Ajuste por moneda nativa: ","un sueldo dolarizado no pierde valor con la inflación local; se "
       "ajusta por la inflación del país de su moneda, de modo que el cociente pesos/dólares sea un "
       "tipo de cambio coherente.")
bullet("Outliers: ","se conservan marcados; sólo se descartan errores de carga evidentes (>50× la "
       "mediana).")

# 6 FEATURE ENGINEERING (NUEVO, justificado)
h("6. Feature engineering",1)
para("El feature engineering convierte el perfil crudo en variables que el modelo puede aprovechar. "
     "Cada transformación responde a una justificación concreta:")
bullet("Tecnologías — excluidas como features: ","el campo multivaluado se evaluó como multi-hot "
       "(top 20), pero se excluyó del modelo final porque su premium es un proxy del contexto y "
       "generaba ruido en perfiles atípicos (detalle en «Roles de alta cardinalidad y tecnologías»). "
       "El análisis descriptivo se conserva en el EDA.")
bullet("Agrupación de roles (698 → top 15 + «Otro»): ","las categorías con pocos casos producen "
       "dummies ruidosas y overfitting; agruparlas concentra la señal.")
bullet("Agrupación de provincias (<100 registros → «Otra»): ","misma justificación que en los roles.")
bullet("One-hot de las categóricas ","(provincia, género, modalidad, tamaño de empresa, rol, cobra "
       "en dólares): los modelos requieren entradas numéricas y el one-hot no impone un orden falso.")
bullet("Ingeniería del target: ","el salario nominal se transformó en salario real (deflación por "
       "IPC para pesos y por US CPI para dólares, base mayo 2026). Justificación: es lo que hace "
       "comparable el objetivo entre ediciones; sin esto el modelo aprendería el paso del tiempo, no "
       "el perfil.")
bullet("Selección de variables: ","se excluyó «seniority» por ser redundante con los años de "
       "experiencia (multicolinealidad), y se excluyeron las variables derivadas del propio salario "
       "(USD, canastas básicas) para evitar fuga de información (data leakage).")
para("Las variables numéricas se estandarizan dentro de un Pipeline de scikit-learn —ajustado sólo "
     "con los datos de entrenamiento— para los modelos lineales; los árboles no lo requieren. Este "
     "encapsulado garantiza que no haya filtración entre train y test.",italic=True)

para("Tratamiento de outliers",bold=True,color=TEAL)
para("Los salarios extremos se marcaron con una bandera (columna es_outlier: por debajo del "
     "percentil 1 o por encima del 99 del salario real) pero NO se eliminaron: son 622 casos (2% "
     "del total). Durante la limpieza sólo se descartaron los errores de carga evidentes —valores "
     "fuera de 50 veces la mediana de su edición, típicamente un sueldo cargado en miles o en la "
     "moneda equivocada—. El modelo predictivo conserva los outliers, para no sesgar el aprendizaje "
     "eliminando salarios altos legítimos (un CTO o un perfil muy senior es un outlier válido); la "
     "bandera queda disponible por si se quisiera excluirlos en un análisis posterior.")

para("Roles de alta cardinalidad y tecnologías",bold=True,color=TEAL)
para("El rol y las tecnologías son campos con muchísimos valores distintos (698 roles y 396 "
     "tecnologías únicas). Se trataron con dos decisiones distintas:")
bullet("Roles — agrupación top-N: ","se conservan los 15 roles más frecuentes y el resto se agrupa "
       "en «Otro». Los 15 principales ya cubren el 90% de los registros; los otros 683 suman apenas "
       "el 10%, con muy pocos casos cada uno.")
bullet("Tecnologías — EXCLUIDAS del modelo: ","como features individuales generaban ruido y "
       "extrapolación en perfiles atípicos (poner «Go» en un helpdesk inflaba la predicción hacia "
       "una combinación que no existe en los datos), y su premium es en realidad un proxy del "
       "contexto laboral (backend/infra, empresas dolarizadas), no un efecto del lenguaje. Quitarlas "
       "casi no afecta el desempeño (el R² baja de 0,30 a 0,27) y da un modelo más robusto. El "
       "análisis descriptivo de tecnologías se conserva en el EDA (hipótesis H3).")
para("Agrupar los roles concentra la señal sin perder capacidad predictiva. La curva de Pareto lo "
     "muestra: unos pocos roles/tecnologías concentran casi todo, y el resto es una cola larga.")
img(PRE/"cobertura_roles_techs.png",6.2,"Distribución de frecuencia de roles y tecnologías (Pareto): "
    "el top-N concentra la señal; la cola larga aporta poco.")

# 7 EDA
h("7. Análisis exploratorio (EDA)",1)
para("Todo se analiza en moneda real de mayo 2026. La mediana del mercado es de $3,19 M (USD 2.264) "
     "mensuales brutos.")
img(PRE/"distribucion_bn.png",6.3,"Figura 1. Distribución de salarios reales (pesos y dólares).")
para("Seniority ordena el salario casi sin solapamiento (mediana junior $1,8 M; senior $4,3 M). CABA "
     "paga +24% que el interior y el trabajo remoto +43% en dólares que el presencial. El premium por "
     "tecnología depende de cuál se use (Go +45%, Rust +48%, Python +13%), no de cuántas.")
img(EDA/"eda_04_tecnologias.png",6.0,"Figura 2. Frecuencia y premium salarial por tecnología.")

# 8 RANDOM FOREST
h("8. Modelado supervisado — Random Forest",1)
para("Variable objetivo: el salario real (entrenado en pesos y, en paralelo, en dólares; como "
     "difieren sólo en una constante, el R² es idéntico). Se compararon cinco técnicas con validación "
     "temporal: se entrena con las ediciones previas y se evalúa sobre la última (2025.2), nunca vista.")
table(["Modelo","R² (test)","RMSE (test)","Overfitting"],[
    ["Regresión lineal simple (sólo experiencia)","0.13","$2,18 M","no"],
    ["Regresión lineal múltiple","0.22","$2,07 M","no"],
    ["Regresión polinómica (grado 2)","0.28","$1,99 M","no"],
    ["Árbol de decisión (prof. 5)","0.20","$2,10 M","no"],
    ["Random Forest (regularizado)","0.27","$2,00 M","no"],
])
para("Por qué elegimos Random Forest",bold=True,color=TEAL)
para("El Random Forest y la regresión polinómica empatan en test (R²≈0,27–0,28). Se eligió el "
     "Random Forest por cuatro razones: (1) captura no linealidades e interacciones sin ingeniería "
     "manual —por ejemplo, el rendimiento decreciente de la experiencia—; (2) es robusto a outliers "
     "y a la mezcla de escalas, y no requiere estandarización; (3) regularizado no presenta "
     "overfitting (train ≈ test); y (4) entrega la importancia de cada variable, aportando "
     "interpretabilidad para la lectura de negocio. Frente a los lineales (que no capturan "
     "interacciones) y al árbol único (inestable), el bosque ofrece el mejor balance entre robustez "
     "e interpretabilidad.")
img(MOD/"modelos_importancia_features.png",5.6,"Figura 3. Importancia de variables (Random Forest).")
para("Las variables más predictivas son la experiencia, la edad, el rol de liderazgo, la provincia "
     "(CABA) y el tamaño de empresa. El techo de ~0,27 es esperable y honesto: cerca de dos tercios "
     "de la variación dependen de factores no encuestados (empresa puntual, negociación). La "
     "validación temporal (R²≈0,27) es muy cercana a la aleatoria (R²≈0,29): el modelo generaliza "
     "bien al presente.")

# 9 KMEANS
h("9. Segmentación — K-Means",1)
para("Con K-Means se agruparon los perfiles sin usar el salario, y luego se observó cuánto gana cada "
     "grupo. El número de segmentos (k=4) se eligió con el método del codo y la silueta. Para "
     "visualizarlos se aplicó PCA: el Componente 1 (48% de la variación) representa los años de "
     "carrera (junior→senior) y el Componente 2 (15%) la antigüedad/rotación en la empresa.")
img(PRE/"clusters_bn.png",6.3,"Figura 4. Mapa de segmentos (PCA) y salario por arquetipo.")
table(["Segmento","%","Perfil","Salario mediano"],[
    ["Junior remoto","49%","28–29 años, 3 exp, remoto, en pesos","$2,75 M"],
    ["Semi-senior dolarizado","21%","31 años, 5 exp, 100% en USD","$2,81 M"],
    ["Senior remoto","21%","~40 años, 15 exp, remoto","$4,43 M"],
    ["Senior +15 corporativo estable","9%","45 años, 18 exp, 14 años en la empresa","$4,01 M"],
])
para("Nota para la defensa — por qué los grupos se ven solapados en el gráfico:",bold=True,color=TEAL)
para("K-Means asigna cada punto a su centroide más cercano en el espacio completo (lo verificamos: "
     "el 100% de los casos). En el scatter se ven solapados porque es una proyección 2D (PCA) que "
     "muestra sólo el 62% de la información —pierde el 38% que separa a los grupos— y porque los "
     "perfiles forman un continuo sin fronteras naturales (la transición junior→senior es gradual). "
     "No es un error del modelo: en las dimensiones reales los segmentos tienen perfiles nítidamente "
     "distintos (ver la tabla de medias por segmento).")

# 10 RESULTADOS
h("10. Resultados y storytelling",1)
para("Validación de las hipótesis:",bold=True)
table(["#","Resultado","Evidencia"],[
    ["H1","Confirmada","Distribuciones casi sin solape; el predictor más fuerte."],
    ["H2","Confirmada","0→3 años: +85%; 10→13 años: +3% (plateau)."],
    ["H3","Parcial","Premium real, pero liderado por Go/Rust, no sólo Python."],
    ["H4","Confirmada","CABA +24% vs interior; remoto +43% en USD."],
    ["H5","Refutada","Los dolarizados tuvieron MÁS volatilidad real (CV 0,23 vs 0,14)."],
    ["H6","Indicio","Correlación 0,70 pero no significativa (n=6)."],
])
para("Brechas salariales y comparativa regional",bold=True,color=TEAL)
img(EDA/"eda_05_geografia.png",6.0,"Figura 5. Salario real mediano por provincia.")
para("La brecha junior → senior es de casi 3× y la de liderazgo es la mayor entre roles (+0,19 en "
     "log-salario). En lo regional, CABA paga un 24% más que el interior, la Patagonia destaca por "
     "el polo energético, y las grandes empresas y el stack backend concentran el premium. Contra "
     "la creencia general, cobrar en dólares no dio un poder de compra más estable: los dolarizados "
     "mostraron más volatilidad real que los pesificados —el dólar fue escudo en 2022–2023 y lastre "
     "en 2024–2026— (H5 refutada).")
para("Comparativa internacional (Stack Overflow)",bold=True,color=TEAL)
img(PRE/"comparacion_mundo.png",6.0,"Figura 6. Salario mediano de developers por país (USD/mes).")
para("Según la encuesta global de Stack Overflow, el developer argentino gana una mediana de USD "
     "3.333 mensuales: por encima de Brasil, Ucrania e India, pero muy por debajo de Europa y de "
     "EE.UU. (USD 12.500). La medición de Sysarmy es más conservadora (USD 2.264) porque Stack "
     "Overflow capta perfiles más internacionales y senior. Replicar el modelo predictivo a nivel "
     "internacional es una línea de trabajo futura: exigiría rehacer el feature engineering, "
     "incorporar el país como variable dominante y ajustar por poder adquisitivo entre países.")

# 11 APP
h("11. La herramienta (aplicación web)",1)
para("Los resultados se entregan en una app Streamlit con tres secciones: un estimador de salario "
     "(Random Forest), un EDA navegable y la segmentación de mercado. Dos versiones:")
bullet("Español / pesos reales — streamlit run app/main.py")
bullet("Inglés / dólares reales — streamlit run app_en/main.py")

# 12 CONCLUSIONES
h("12. Conclusiones",1)
para("Principales hallazgos:",bold=True)
bullet("El perfil explica ~27% del sueldo; experiencia y rol de liderazgo son lo decisivo.")
bullet("Dolarizarse no garantizó ni mayor premium ni mayor estabilidad.")
bullet("CABA, las grandes empresas y el stack backend concentran los mejores sueldos.")
para("Valor estratégico para la gerencia:",bold=True)
bullet("Benchmarking salarial objetivo y segmentado por arquetipo.")
bullet("Lectura del costo del talento frente al contexto macroeconómico.")
bullet("Decisiones de oferta y retención basadas en datos, con una herramienta reutilizable.")

# ANEXO GUION
doc.add_page_break(); h("Anexo — Guion de defensa oral",1)
guion=[("Portada","Presentarse y enunciar el tema."),
       ("Equipo y estrategia","Roles del grupo y metodología (CRISP-DM)."),
       ("Dominio","El problema: info fragmentada y no comparable. Hipótesis central."),
       ("Hipótesis","Enunciar H1–H6."),
       ("Propuesta y valor","Qué construimos y para qué le sirve a la gerencia."),
       ("Fuentes","Las 9 fuentes integradas."),
       ("Pipeline","El flujo de 4 pasos."),
       ("Limpieza","Foco: armonización, faltantes, moneda nativa, outliers."),
       ("Feature engineering","Foco: agrupaciones, target real, anti-leakage, techs excluidas."),
       ("EDA distribución","Trabajamos en moneda real; mediana $3,19M / USD 2.264."),
       ("EDA seniority/geo","H1 y H4."),
       ("EDA tecnologías","H3: premium por stack."),
       ("Dispersión / ejes","K-Means y la lectura del PCA; 4 arquetipos."),
       ("Variable target","Salario real, no nominal; dos unidades."),
       ("Comparación de modelos","5 técnicas; gana RF."),
       ("Por qué RF","Las 5 razones."),
       ("Dificultades","Lucirse con las decisiones metodológicas."),
       ("Resultados","H1–H6, brechas, regional y dolarización (H5 refutada)."),
       ("Demo","Mostrar la app en vivo."),
       ("Conclusiones","Hallazgos y valor estratégico.")]
for t,d in guion:
    p=doc.add_paragraph(style="List Number"); r=p.add_run(f"{t}: "); r.bold=True; p.add_run(d)

doc.save(str(OUT))
print("DOCX:", OUT, "| Headings:", sum(1 for p in doc.paragraphs if p.style.name=="Heading 1"))
